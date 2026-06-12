import os
# 必须在最开头设置镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

import re
import glob
import pickle
from docx import Document
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_deepseek import ChatDeepSeek
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableParallel
from dotenv import load_dotenv

load_dotenv()

# ------------------------------
# 1. 读取 docx 文件并清洗文本
# ------------------------------
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
            text = re.sub(r'\{.*?\}', '', text)
            text = re.sub(r'[ ]+', ' ', text)
            full_text.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    full_text.append(cell_text)
    return "\n".join(full_text)

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8") as f:
        return f.read()

def extract_text_from_pdf(pdf_path):
    from pypdf import PdfReader
    reader = PdfReader(pdf_path)
    text = "\n".join([page.extract_text() or "" for page in reader.pages])
    return text

# ------------------------------
# 2. 加载 data 文件夹下所有支持的文件
# ------------------------------
data_dir = "data"
all_docs = []

extensions = ["*.docx", "*.txt", "*.pdf"]
for ext in extensions:
    for file_path in glob.glob(os.path.join(data_dir, ext)):
        print(f"正在加载文件: {file_path}")
        if ext == "*.docx":
            content = extract_text_from_docx(file_path)
        elif ext == "*.txt":
            content = extract_text_from_txt(file_path)
        elif ext == "*.pdf":
            content = extract_text_from_pdf(file_path)
        else:
            continue
        doc = LCDocument(
            page_content=content,
            metadata={"source": os.path.basename(file_path)}
        )
        all_docs.append(doc)

print(f"共加载 {len(all_docs)} 个文档")

# ------------------------------
# 3. 文本分割
# ------------------------------
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=600,
    chunk_overlap=100,
    separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]
)
documents = text_splitter.split_documents(all_docs)
print(f"切分完成，共 {len(documents)} 个文本片段")

# ------------------------------
# 4. 向量化并存储到 ChromaDB（在线模型，镜像加速）
# ------------------------------
embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-small-zh-v1.5")
vectorstore = Chroma.from_documents(
    documents=documents,
    embedding=embeddings,
    persist_directory="./chroma_db"
)
print("向量库构建完成，已保存到 ./chroma_db")

# 保存所有文本片段用于 BM25
all_texts = [doc.page_content for doc in documents]
with open("all_texts.pkl", "wb") as f:
    pickle.dump(all_texts, f)
print("已保存 all_texts.pkl 用于 BM25")

# ------------------------------
# 5. 构建 RAG 链（测试用）
# ------------------------------
llm = ChatDeepSeek(
    model=os.getenv("DEEPSEEK_MODEL", "deepseek-chat"),
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1"),
    temperature=0.3,
)

retriever = vectorstore.as_retriever(search_kwargs={"k": 4})

prompt_template = """你是一个专业的制药质量与中药鉴定助手。请根据以下从知识库中检索到的片段回答用户问题。

【知识库片段】
{context}

【用户问题】
{question}

要求：
1. 只使用上述片段中的信息，不要编造。
2. 如果片段中没有足够信息，请直接说“知识库中暂未找到相关信息”。
3. 回答结束后，请用【来源：xxx】列出主要参考的文档名称。

回答："""
prompt = ChatPromptTemplate.from_messages([("system", prompt_template), ("human", "{question}")])

def format_docs(docs):
    return "\n\n---\n\n".join([
        f"【来源：{doc.metadata.get('source', '未知')}】\n{doc.page_content}"
        for doc in docs
    ])

rag_chain = (
    RunnableParallel(
        context=retriever | format_docs,
        question=RunnablePassthrough()
    )
    | prompt
    | llm
    | StrOutputParser()
)

if __name__ == "__main__":
    questions = [
        "药典2025年版一部收载了多少个品种？",
        "人参的性状特征是什么？",
        "GMP中关于产品召回的要求有哪些？"
    ]
    for q in questions:
        print("\n问题：", q)
        answer = rag_chain.invoke(q)
        print("回答：", answer)
        print("-" * 50)
